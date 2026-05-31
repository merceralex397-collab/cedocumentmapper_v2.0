from __future__ import annotations

import io
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from cedocumentmapper_v2.domain.models import ExtractedRecord, FieldKey
from cedocumentmapper_v2.exporters.base import Exporter


def format_date_for_rjs(value: str) -> str:
    """Format DD/MM/YYYY date to 'Day Month Year' (e.g. 15th April 2026)."""
    raw = (value or "").strip()
    try:
        dt = datetime.strptime(raw, "%d/%m/%Y")
        day = dt.day
        if 10 <= day % 100 <= 20:
            suffix = "th"
        else:
            suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
        return dt.strftime(f"{day}{suffix} %B %Y")
    except Exception:
        return raw


def split_address_lines(value: str) -> list[str]:
    text = (value or "").strip()
    if not text:
        return []
    if "\n" in text:
        parts = [part.strip() for part in text.splitlines() if part.strip()]
    else:
        parts = [part.strip() for part in re.split(r",\s*", text) if part.strip()]
    return parts[:6]


def normalise_rjs_address_block(value: str) -> list[str]:
    parts = split_address_lines(value)
    if not parts:
        return ["", "", ""]
    if len(parts) == 1:
        return [parts[0], "", ""]
    if len(parts) == 2:
        return [parts[0], parts[1], ""]
    if len(parts) == 3:
        return parts
    return [parts[0], parts[1], parts[-1]]


class RJSDocxExporter(Exporter):
    def export(self, record: ExtractedRecord) -> bytes:
        """Build the RJS letter and return as docx bytes."""
        doc = Document()
        
        # Configure default Normal style
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(12)

        # Margins
        section = doc.sections[0]
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.0)

        # Helper
        def add_paragraph(text: str = "", bold: bool = False, center: bool = False, space_after: int = 0) -> None:
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_after = Pt(space_after)
            pf.space_before = Pt(0)
            pf.line_spacing = 1.0
            if center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.name = "Arial"
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Arial")
            run.font.size = Pt(12)
            run.bold = bold

        # Top Address Block
        for line in [
            "Collision Engineers Ltd",
            "2 Castle Buildings",
            "147 - 149 Telegraph Road",
            "Heswall",
            "Wirral",
            "United Kingdom",
            "CH60 7SE",
        ]:
            add_paragraph(line)
        add_paragraph(space_after=10)

        # Header values
        ref_val = record.fields.get(FieldKey.REFERENCE)
        ref_str = ref_val.value if ref_val else ""
        add_paragraph("Your Reference:")
        add_paragraph(f"Our Reference: {ref_str.strip()}")
        add_paragraph("Fee earner: Keeley Garner")
        add_paragraph("Direct dial: 01516650836")
        add_paragraph("Email: k.garner@robertjameslaw.co.uk")

        inst_date_val = record.fields.get(FieldKey.INSTRUCTION_DATE)
        inst_date_str = inst_date_val.value if inst_date_val else ""
        add_paragraph(format_date_for_rjs(inst_date_str), space_after=22)

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(14)
        run = p.add_run("URGENT VEHICLE INSPECTION REQUIRED")
        run.font.name = "Arial"
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(14)
        run.bold = True

        add_paragraph("Dear Sirs", space_after=14)

        # Details Block
        def add_label_value(label: str, value: str) -> None:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r1 = p.add_run(f"{label}:    ")
            r1.font.name = "Arial"
            r1._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Arial")
            r1.font.size = Pt(12)
            r1.bold = True
            
            r2 = p.add_run(value)
            r2.font.name = "Arial"
            r2._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Arial")
            r2.font.size = Pt(12)
            r2.bold = True

        claimant_val = record.fields.get(FieldKey.CLAIMANT_NAME)
        claimant_str = claimant_val.value if claimant_val else ""
        
        inc_date_val = record.fields.get(FieldKey.INCIDENT_DATE)
        inc_date_str = inc_date_val.value if inc_date_val else ""

        vrm_val = record.fields.get(FieldKey.VRM)
        vrm_str = vrm_val.value if vrm_val else ""

        model_val = record.fields.get(FieldKey.VEHICLE_MODEL)
        model_str = model_val.value if model_val else ""

        add_label_value("Our Client", claimant_str.strip())
        add_label_value("Accident", format_date_for_rjs(inc_date_str.strip()))
        add_label_value("Client vehicle registration", vrm_str.strip())
        add_label_value("Client vehicle make", "")
        add_label_value("Client vehicle model", model_str.strip())

        add_paragraph(space_after=16)

        # Narratives
        claimant_name = claimant_str.strip() or "the above named client"
        formatted_inc_date = format_date_for_rjs(inc_date_str.strip())

        narrative_1 = (
            f"I act on behalf of my above named client in the recovery of damages resulting from an accident "
            f"which occurred on {formatted_inc_date}."
        )
        narrative_2 = (
            "Please arrange an inspection of my client’s vehicle as soon as possible and provide a report "
            "detailing the damage sustained, costs of repair or cost of replacement if beyond repair."
        )
        narrative_3 = (
            f"I have advised my client of your instruction. Please make arrangements for the inspection with my "
            f"client. {claimant_name} is available at:"
        )

        add_paragraph(narrative_1, space_after=14)
        add_paragraph(narrative_2, space_after=14)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)
        run_n3 = p.add_run(narrative_3)
        run_n3.font.name = "Arial"
        run_n3._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Arial")
        run_n3.font.size = Pt(12)
        run_n3.bold = True

        # Address
        addr_val = record.fields.get(FieldKey.INSPECTION_ADDRESS)
        addr_str = addr_val.value if addr_val else ""
        address_lines = normalise_rjs_address_block(addr_str)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r1 = p.add_run("Address:    ")
        r1.font.name = "Arial"
        r1._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Arial")
        r1.font.size = Pt(12)
        
        r2 = p.add_run(address_lines[0])
        r2.font.name = "Arial"
        r2._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Arial")
        r2.font.size = Pt(12)

        add_paragraph(address_lines[1])
        add_paragraph(address_lines[2])

        # RJS spacing
        for _ in range(4):
            add_paragraph("")

        add_paragraph("Mobile Tel:")
        add_paragraph(space_after=18)

        boilerplate = (
            "I can confirm that in accordance with the Civil Procedure Rules I have notified the third party of "
            "your involvement in this matter on your behalf and confirmed that I will copy them with your report "
            "once available. Once they are in receipt of your report they may choose to contact you direct with "
            "questions concerning my clients losses. I would be grateful if you could ensure that I receive copies "
            "of any such correspondence along with your replies."
        )
        fees = (
            "Finally having regard to your reasonable fees, I hereby confirm that this firm will be responsible for "
            "the same in accordance with our agreement."
        )

        add_paragraph(boilerplate, space_after=14)
        add_paragraph(fees, space_after=18)
        add_paragraph("Yours faithfully", space_after=40)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run_kg = p.add_run("Keeley Garner")
        run_kg.font.name = "Arial"
        run_kg._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Arial")
        run_kg.font.size = Pt(12)
        run_kg.bold = True

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run_rjl = p.add_run("Robert James Solicitors")
        run_rjl.font.name = "Arial"
        run_rjl._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Arial")
        run_rjl.font.size = Pt(12)
        run_rjl.bold = True

        # Save to BytesIO
        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()
