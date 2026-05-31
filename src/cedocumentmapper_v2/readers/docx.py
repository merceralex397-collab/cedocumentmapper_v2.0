from __future__ import annotations

import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path
from docx import Document

from cedocumentmapper_v2.domain.models import (
    DocumentModel,
    DocumentPage,
    DocumentLine,
)
from cedocumentmapper_v2.readers.base import DocumentReader
from cedocumentmapper_v2.readers.errors import ReaderError


class DocxDocumentReader(DocumentReader):
    supported_extensions: frozenset[str] = frozenset([".docx"])

    def read(self, path: Path) -> DocumentModel:
        if not path.exists():
            raise ReaderError(f"File not found: {path}")

        notes: list[str] = []
        lines_list = []
        line_idx_counter = 0

        # Helper to append unique clean lines
        seen_lower = set()

        def add_line(text: str, group_id: str | None = None) -> None:
            nonlocal line_idx_counter
            cleaned = text.replace("\r", " ").replace("\t", " ").strip()
            if not cleaned:
                return
            lowered = cleaned.lower()
            if lowered not in seen_lower:
                lines_list.append(
                    DocumentLine(
                        text=cleaned,
                        page_index=0,
                        line_index=line_idx_counter,
                        block_id=group_id,
                        confidence=1.0,
                    )
                )
                seen_lower.add(lowered)
                line_idx_counter += 1

        try:
            # 1. Read headers and footers from zip structure
            headers, footers = self._extract_docx_header_footer_lines(path)
            for line in headers:
                add_line(line, "header")

            # 2. Read body paragraphs using python-docx
            doc = Document(str(path))
            for para in doc.paragraphs:
                if para.text.strip():
                    add_line(para.text, "body")

            # 3. Read body tables using python-docx
            for table_idx, table in enumerate(doc.tables):
                for row in table.rows:
                    row_parts = [(cell.text or "").strip() for cell in row.cells]
                    if any(row_parts):
                        # Combine cell values with pipes
                        combined = " | ".join(p for p in row_parts if p)
                        add_line(combined, f"table_{table_idx}")

            # 4. Read textboxes from zip structure
            textbox_lines = self._extract_docx_textbox_lines(path)
            for line in textbox_lines:
                add_line(line, "textbox")

            # 5. Read footers
            for line in footers:
                add_line(line, "footer")

        except Exception as exc:
            raise ReaderError(f"Could not read DOCX: {exc}") from exc

        # Create single page for unpaginated Word document
        page = DocumentPage(
            page_index=0,
            lines=tuple(lines_list),
        )

        combined_text = "\n".join(l.text for l in lines_list)

        return DocumentModel(
            source_path=path,
            source_type="docx",
            pages=(page,),
            plain_text=combined_text,
            reader_notes=tuple(notes),
            metadata={
                "raw_text": combined_text,
                "raw_lines": combined_text.replace("\r\n", "\n").replace("\r", "\n").split("\n") if combined_text else [],
                "line_groups": sorted({line.block_id for line in lines_list if line.block_id}),
            },
        )

    def _extract_docx_textbox_lines(self, path: Path) -> list[str]:
        textbox_lines = []
        try:
            with zipfile.ZipFile(str(path)) as zf:
                xml_names = [
                    name for name in zf.namelist()
                    if name.startswith("word/") and name.endswith(".xml")
                ]
                for name in xml_names:
                    try:
                        root = ET.fromstring(zf.read(name))
                    except Exception:
                        continue

                    for node in root.iter():
                        tag_name = node.tag.rsplit("}", 1)[-1].lower()
                        if tag_name not in {"txbxcontent", "textbox"}:
                            continue

                        collected = []
                        for child in node.iter():
                            child_tag = child.tag.rsplit("}", 1)[-1].lower()
                            if child_tag == "t":
                                txt = (child.text or "").strip()
                                if txt:
                                    collected.append(txt)
                            elif child_tag in {"p", "br", "cr"} and collected:
                                textbox_lines.append(" ".join(collected).strip())
                                collected = []
                        if collected:
                            textbox_lines.append(" ".join(collected).strip())
        except Exception:
            return []

        cleaned = []
        seen = set()
        for line in textbox_lines:
            line_cleaned = line.replace("\r", " ").replace("\t", " ").strip(" :\n")
            if line_cleaned and line_cleaned.lower() not in seen:
                cleaned.append(line_cleaned)
                seen.add(line_cleaned.lower())
        return cleaned

    def _extract_docx_header_footer_lines(self, path: Path) -> tuple[list[str], list[str]]:
        def dedupe(lines: list[str]) -> list[str]:
            output = []
            seen = set()
            for line in lines:
                key = line.lower()
                if key not in seen:
                    output.append(line)
                    seen.add(key)
            return output

        def extract_lines_from_parts(zf: zipfile.ZipFile, part_prefix: str) -> list[str]:
            collected = []
            xml_names = [
                name for name in zf.namelist()
                if name.startswith("word/") and name.endswith(".xml") and Path(name).name.lower().startswith(part_prefix)
            ]
            for name in xml_names:
                try:
                    root = ET.fromstring(zf.read(name))
                except Exception:
                    continue

                line_parts = []
                for node in root.iter():
                    tag_name = node.tag.rsplit("}", 1)[-1].lower()
                    if tag_name == "t":
                        txt = node.text or ""
                        if txt:
                            line_parts.append(txt)
                    elif tag_name == "tab":
                        line_parts.append(" ")
                    elif tag_name in {"br", "cr", "p"}:
                        val = "".join(line_parts).replace("\r", " ").replace("\t", " ").strip(" :\n")
                        if val:
                            collected.append(val)
                        line_parts = []
                val = "".join(line_parts).replace("\r", " ").replace("\t", " ").strip(" :\n")
                if val:
                    collected.append(val)
            return dedupe(collected)

        try:
            with zipfile.ZipFile(str(path)) as zf:
                headers = extract_lines_from_parts(zf, "header")
                footers = extract_lines_from_parts(zf, "footer")
                # Try core title
                try:
                    core_root = ET.fromstring(zf.read("docProps/core.xml"))
                    title_node = None
                    for node in core_root.iter():
                        node_text = node.text
                        if node.tag.rsplit("}", 1)[-1].lower() == "title" and node_text and node_text.strip():
                            title_node = node_text.replace("\r", " ").replace("\t", " ").strip(" :\n")
                            break
                    if title_node:
                        headers = dedupe([title_node] + headers)
                except Exception:
                    pass
                return headers, footers
        except Exception:
            return [], []
